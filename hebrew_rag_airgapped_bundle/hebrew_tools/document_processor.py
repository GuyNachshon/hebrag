# Hebrew Document Processing Tools
from agno.tools import Tool
from typing import List, Dict, Optional, Any, Union
import fitz  # PyMuPDF
import pdfplumber
import logging
from pathlib import Path
import re
import json
from PIL import Image
import numpy as np

try:
    import paddleocr
    PADDLE_OCR_AVAILABLE = True
except ImportError:
    PADDLE_OCR_AVAILABLE = False
    logging.warning("PaddleOCR not available. OCR functionality will be limited.")

try:
    from detectron2.engine import DefaultPredictor
    from detectron2.config import get_cfg
    DETECTRON2_AVAILABLE = True
except ImportError:
    DETECTRON2_AVAILABLE = False
    logging.warning("Detectron2 not available. Layout analysis will be limited.")

class HebrewDocumentProcessor(Tool):
    """
    Hebrew Document Processor for Agno agents.
    Processes Hebrew documents extracting text and visual elements with spatial context.
    """
    
    def __init__(self, ocr_enabled: bool = True, layout_analysis: bool = True, visual_extraction: bool = True):
        super().__init__(
            name="hebrew_document_processor",
            description="Process Hebrew documents extracting text and visual elements with spatial context"
        )
        self.ocr_enabled = ocr_enabled and PADDLE_OCR_AVAILABLE
        self.layout_analysis = layout_analysis and DETECTRON2_AVAILABLE
        self.visual_extraction = visual_extraction
        self.logger = logging.getLogger(__name__)
        self.setup_models()
    
    def setup_models(self):
        """Initialize Hebrew OCR and layout analysis models"""
        try:
            if self.ocr_enabled:
                self.ocr = paddleocr.PaddleOCR(
                    use_angle_cls=True,
                    lang='he',  # Hebrew
                    use_gpu=False,  # Set to True if GPU available
                    show_log=False
                )
                self.logger.info("Hebrew OCR model initialized")
            
            if self.layout_analysis:
                cfg = self.setup_detectron_config()
                self.table_detector = DefaultPredictor(cfg)
                self.logger.info("Layout analysis model initialized")
                
        except Exception as e:
            self.logger.error(f"Error setting up models: {e}")
            self.ocr_enabled = False
            self.layout_analysis = False
    
    def setup_detectron_config(self):
        """Setup Detectron2 configuration for table detection"""
        cfg = get_cfg()
        # Use a basic COCO detection model as fallback
        cfg.MODEL.DEVICE = "cpu"  # Use CPU for air-gapped deployment
        return cfg
    
    def run(self, document_path: str) -> Dict[str, Any]:
        """
        Process a Hebrew document and return structured content
        
        Args:
            document_path: Path to PDF, DOCX, or other document
            
        Returns:
            Dict containing processed content with spatial relationships
        """
        try:
            document_path = Path(document_path)
            if not document_path.exists():
                return {
                    "status": "error",
                    "error": f"Document not found: {document_path}",
                    "document_path": str(document_path)
                }
            
            # Extract layout-aware content
            elements = self.extract_layout_elements(str(document_path))
            
            # Process Hebrew text
            hebrew_content = self.process_hebrew_text(elements)
            
            # Extract and analyze visual elements
            visual_elements = self.extract_visual_elements(str(document_path))
            
            # Create contextual chunks preserving relationships
            contextual_chunks = self.create_contextual_chunks(
                hebrew_content, visual_elements
            )
            
            return {
                "status": "success",
                "document_path": str(document_path),
                "chunks": contextual_chunks,
                "visual_elements_count": len(visual_elements),
                "hebrew_text_blocks": len(hebrew_content),
                "total_chunks": len(contextual_chunks)
            }
            
        except Exception as e:
            self.logger.error(f"Error processing document {document_path}: {e}")
            return {
                "status": "error", 
                "error": str(e),
                "document_path": str(document_path)
            }
    
    def extract_layout_elements(self, document_path: str) -> List[Dict]:
        """Extract layout-aware elements from document"""
        elements = []
        
        try:
            if document_path.lower().endswith('.pdf'):
                elements = self._extract_pdf_elements(document_path)
            elif document_path.lower().endswith(('.docx', '.doc')):
                elements = self._extract_docx_elements(document_path)
            else:
                # Try to process as PDF
                elements = self._extract_pdf_elements(document_path)
                
        except Exception as e:
            self.logger.error(f"Error extracting layout elements: {e}")
            
        return elements
    
    def _extract_pdf_elements(self, pdf_path: str) -> List[Dict]:
        """Extract elements from PDF using pdfplumber for better layout preservation"""
        elements = []
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text with coordinates
                    chars = page.chars
                    if chars:
                        # Group characters into words and lines
                        text_elements = self._group_text_elements(chars, page_num)
                        elements.extend(text_elements)
                    
                    # Extract tables
                    tables = page.find_tables()
                    for table in tables:
                        table_element = {
                            'type': 'table',
                            'content': table.extract(),
                            'bbox': table.bbox,
                            'page_number': page_num
                        }
                        elements.append(table_element)
                    
                    # Extract images
                    if hasattr(page, 'images'):
                        for img in page.images:
                            img_element = {
                                'type': 'image',
                                'bbox': (img['x0'], img['top'], img['x1'], img['bottom']),
                                'page_number': page_num
                            }
                            elements.append(img_element)
                            
        except Exception as e:
            self.logger.error(f"Error extracting PDF elements: {e}")
            
        return elements
    
    def _extract_docx_elements(self, docx_path: str) -> List[Dict]:
        """Extract elements from DOCX file"""
        elements = []
        
        try:
            from docx import Document
            doc = Document(docx_path)
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    element = {
                        'type': 'text',
                        'content': paragraph.text,
                        'paragraph_number': para_num,
                        'style': paragraph.style.name if paragraph.style else None
                    }
                    elements.append(element)
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                table_element = {
                    'type': 'table',
                    'content': table_data,
                    'table_number': table_num
                }
                elements.append(table_element)
                
        except Exception as e:
            self.logger.error(f"Error extracting DOCX elements: {e}")
            
        return elements
    
    def _group_text_elements(self, chars: List[Dict], page_num: int) -> List[Dict]:
        """Group character elements into meaningful text blocks"""
        elements = []
        
        if not chars:
            return elements
        
        # Sort characters by position (top to bottom, right to left for Hebrew)
        sorted_chars = sorted(chars, key=lambda x: (x['top'], -x['x0']))
        
        # Group into lines and words
        lines = []
        current_line = []
        current_y = sorted_chars[0]['top']
        
        for char in sorted_chars:
            # If character is on a significantly different y-coordinate, start new line
            if abs(char['top'] - current_y) > 5:  # 5 pixel tolerance
                if current_line:
                    lines.append(current_line)
                current_line = [char]
                current_y = char['top']
            else:
                current_line.append(char)
        
        if current_line:
            lines.append(current_line)
        
        # Convert lines to text elements
        for line_num, line in enumerate(lines):
            if line:
                text = ''.join([char['text'] for char in line])
                if text.strip():
                    element = {
                        'type': 'text',
                        'content': text.strip(),
                        'bbox': (
                            min(char['x0'] for char in line),
                            min(char['top'] for char in line),
                            max(char['x1'] for char in line),
                            max(char['bottom'] for char in line)
                        ),
                        'page_number': page_num,
                        'line_number': line_num
                    }
                    elements.append(element)
        
        return elements
    
    def process_hebrew_text(self, elements: List[Dict]) -> List[Dict]:
        """Process and enhance Hebrew text elements"""
        hebrew_elements = []
        
        for element in elements:
            if element.get('type') == 'text':
                content = element.get('content', '')
                if self.contains_hebrew(content):
                    # Process Hebrew text
                    processed_element = element.copy()
                    processed_element['content'] = self.normalize_hebrew_text(content)
                    processed_element['is_hebrew'] = True
                    processed_element['text_direction'] = 'rtl'
                    hebrew_elements.append(processed_element)
                elif content.strip():  # Non-Hebrew text
                    processed_element = element.copy()
                    processed_element['is_hebrew'] = False
                    processed_element['text_direction'] = 'ltr'
                    hebrew_elements.append(processed_element)
        
        return hebrew_elements
    
    def contains_hebrew(self, text: str) -> bool:
        """Check if text contains Hebrew characters"""
        hebrew_chars = set(range(0x0590, 0x05FF))  # Hebrew Unicode block
        return any(ord(char) in hebrew_chars for char in text)
    
    def normalize_hebrew_text(self, text: str) -> str:
        """Normalize Hebrew text (remove extra spaces, fix direction, etc.)"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Basic Hebrew text cleaning
        text = text.replace('\u200e', '')  # Remove LTR mark
        text = text.replace('\u200f', '')  # Remove RTL mark
        
        return text
    
    def extract_visual_elements(self, document_path: str) -> List[Dict]:
        """Extract and analyze visual elements"""
        visual_elements = []
        
        if not self.visual_extraction:
            return visual_elements
        
        try:
            if document_path.lower().endswith('.pdf'):
                visual_elements = self._extract_pdf_visuals(document_path)
        except Exception as e:
            self.logger.error(f"Error extracting visual elements: {e}")
        
        return visual_elements
    
    def _extract_pdf_visuals(self, pdf_path: str) -> List[Dict]:
        """Extract visual elements from PDF"""
        visuals = []
        
        try:
            doc = fitz.open(pdf_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract images
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    
                    visual_element = {
                        'type': 'image',
                        'page_number': page_num,
                        'image_index': img_index,
                        'bbox': page.get_image_bbox(img),
                        'size': (pix.width, pix.height)
                    }
                    visuals.append(visual_element)
                    pix = None  # Free memory
                
            doc.close()
        except Exception as e:
            self.logger.error(f"Error extracting PDF visuals: {e}")
        
        return visuals
    
    def create_contextual_chunks(self, text_content: List[Dict], visual_elements: List[Dict]) -> List[Dict]:
        """Create chunks that preserve text-visual relationships"""
        chunks = []
        
        # Group elements by page
        pages = {}
        for element in text_content + visual_elements:
            page_num = element.get('page_number', 0)
            if page_num not in pages:
                pages[page_num] = {'text': [], 'visual': []}
            
            if element.get('type') in ['text']:
                pages[page_num]['text'].append(element)
            else:
                pages[page_num]['visual'].append(element)
        
        # Process each page
        for page_num, page_elements in pages.items():
            page_chunks = self._create_page_chunks(page_num, page_elements)
            chunks.extend(page_chunks)
        
        return chunks
    
    def _create_page_chunks(self, page_num: int, page_elements: Dict) -> List[Dict]:
        """Create contextual chunks for a single page"""
        chunks = []
        text_elements = page_elements.get('text', [])
        visual_elements = page_elements.get('visual', [])
        
        # Create text chunks
        for i, text_element in enumerate(text_elements):
            chunk = {
                'chunk_id': f"page_{page_num}_text_{i}",
                'type': 'text',
                'content': text_element.get('content', ''),
                'page_number': page_num,
                'bbox': text_element.get('bbox'),
                'is_hebrew': text_element.get('is_hebrew', False),
                'text_direction': text_element.get('text_direction', 'ltr'),
                'nearby_visuals': self._find_nearby_visuals(text_element, visual_elements)
            }
            chunks.append(chunk)
        
        # Create visual chunks with surrounding text context
        for i, visual_element in enumerate(visual_elements):
            context_text = self._get_visual_context(visual_element, text_elements)
            
            chunk = {
                'chunk_id': f"page_{page_num}_visual_{i}",
                'type': 'visual',
                'visual_type': visual_element.get('type', 'unknown'),
                'page_number': page_num,
                'bbox': visual_element.get('bbox'),
                'context_before': context_text.get('before', ''),
                'context_after': context_text.get('after', ''),
                'full_context': f"{context_text.get('before', '')} [VISUAL_ELEMENT] {context_text.get('after', '')}".strip()
            }
            
            # Add Hebrew description if OCR is available
            if self.ocr_enabled and visual_element.get('type') == 'image':
                description = self.describe_visual_in_hebrew(visual_element)
                chunk['hebrew_description'] = description
            
            chunks.append(chunk)
        
        return chunks
    
    def _find_nearby_visuals(self, text_element: Dict, visual_elements: List[Dict]) -> List[Dict]:
        """Find visual elements near a text element"""
        nearby = []
        text_bbox = text_element.get('bbox')
        
        if not text_bbox:
            return nearby
        
        for visual in visual_elements:
            visual_bbox = visual.get('bbox')
            if visual_bbox and self._are_elements_nearby(text_bbox, visual_bbox):
                nearby.append({
                    'type': visual.get('type'),
                    'distance': self._calculate_distance(text_bbox, visual_bbox)
                })
        
        return sorted(nearby, key=lambda x: x['distance'])
    
    def _get_visual_context(self, visual_element: Dict, text_elements: List[Dict]) -> Dict[str, str]:
        """Get text context around a visual element"""
        context = {'before': '', 'after': ''}
        visual_bbox = visual_element.get('bbox')
        
        if not visual_bbox:
            return context
        
        # Find text before and after the visual element
        before_texts = []
        after_texts = []
        
        for text_elem in text_elements:
            text_bbox = text_elem.get('bbox')
            if text_bbox:
                # Check if text is above (before) the visual
                if text_bbox[3] < visual_bbox[1]:  # text bottom < visual top
                    before_texts.append((text_elem, self._calculate_distance(text_bbox, visual_bbox)))
                # Check if text is below (after) the visual
                elif text_bbox[1] > visual_bbox[3]:  # text top > visual bottom
                    after_texts.append((text_elem, self._calculate_distance(text_bbox, visual_bbox)))
        
        # Get closest texts
        if before_texts:
            before_texts.sort(key=lambda x: x[1])
            context['before'] = ' '.join([t[0].get('content', '') for t in before_texts[:2]])
        
        if after_texts:
            after_texts.sort(key=lambda x: x[1])
            context['after'] = ' '.join([t[0].get('content', '') for t in after_texts[:2]])
        
        return context
    
    def _are_elements_nearby(self, bbox1: tuple, bbox2: tuple, threshold: float = 100.0) -> bool:
        """Check if two elements are nearby based on their bounding boxes"""
        distance = self._calculate_distance(bbox1, bbox2)
        return distance < threshold
    
    def _calculate_distance(self, bbox1: tuple, bbox2: tuple) -> float:
        """Calculate distance between two bounding boxes"""
        # Center points
        center1 = ((bbox1[0] + bbox1[2]) / 2, (bbox1[1] + bbox1[3]) / 2)
        center2 = ((bbox2[0] + bbox2[2]) / 2, (bbox2[1] + bbox2[3]) / 2)
        
        # Euclidean distance
        return ((center1[0] - center2[0]) ** 2 + (center1[1] - center2[1]) ** 2) ** 0.5
    
    def describe_visual_in_hebrew(self, visual_element: Dict) -> str:
        """Generate Hebrew description of visual element using OCR if available"""
        if not self.ocr_enabled:
            return "אלמנט ויזואלי"  # "Visual element"
        
        try:
            # This is a placeholder - in a real implementation, you would
            # extract the actual image and run OCR on it
            visual_type = visual_element.get('type', 'unknown')
            
            if visual_type == 'table':
                return "טבלה המכילה נתונים"  # "Table containing data"
            elif visual_type == 'image':
                return "תמונה או איור"  # "Image or illustration"
            elif visual_type == 'chart':
                return "תרשים או גרף"  # "Chart or graph"
            else:
                return "אלמנט ויזואלי"  # "Visual element"
                
        except Exception as e:
            self.logger.error(f"Error describing visual element: {e}")
            return "אלמנט ויזואלי"
    
    def get_preceding_context(self, text_content: List[Dict], element: Dict) -> str:
        """Get text context preceding an element"""
        # Simple implementation - in practice, you'd use spatial analysis
        element_idx = text_content.index(element) if element in text_content else -1
        if element_idx > 0:
            return text_content[element_idx - 1].get('content', '')
        return ''
    
    def get_following_context(self, text_content: List[Dict], element: Dict) -> str:
        """Get text context following an element"""
        # Simple implementation - in practice, you'd use spatial analysis
        element_idx = text_content.index(element) if element in text_content else -1
        if element_idx >= 0 and element_idx < len(text_content) - 1:
            return text_content[element_idx + 1].get('content', '')
        return ''
    
    def build_full_context(self, before: str, visual_description: str, after: str) -> str:
        """Build full context string"""
        parts = [part.strip() for part in [before, visual_description, after] if part.strip()]
        return ' '.join(parts)


class LayoutAnalyzer(Tool):
    """
    Layout Analyzer tool for preserving document structure context.
    """
    
    def __init__(self, preserve_context: bool = True, hebrew_aware: bool = True):
        super().__init__(
            name="layout_analyzer",
            description="Analyze document layout and preserve spatial context for Hebrew documents"
        )
        self.preserve_context = preserve_context
        self.hebrew_aware = hebrew_aware
        self.logger = logging.getLogger(__name__)
    
    def run(self, elements: List[Dict], document_info: Dict) -> Dict[str, Any]:
        """
        Analyze layout of document elements
        
        Args:
            elements: List of document elements with spatial information
            document_info: Document metadata
            
        Returns:
            Dict containing layout analysis results
        """
        try:
            # Analyze reading order (important for Hebrew RTL text)
            reading_order = self.analyze_reading_order(elements)
            
            # Identify columns and sections
            layout_structure = self.identify_layout_structure(elements)
            
            # Find relationships between elements
            element_relationships = self.find_element_relationships(elements)
            
            return {
                "status": "success",
                "reading_order": reading_order,
                "layout_structure": layout_structure,
                "element_relationships": element_relationships,
                "total_elements": len(elements)
            }
            
        except Exception as e:
            self.logger.error(f"Error analyzing layout: {e}")
            return {
                "status": "error",
                "error": str(e),
                "elements_count": len(elements)
            }
    
    def analyze_reading_order(self, elements: List[Dict]) -> List[str]:
        """Analyze proper reading order for Hebrew documents"""
        reading_order = []
        
        # Group by page
        pages = {}
        for element in elements:
            page_num = element.get('page_number', 0)
            if page_num not in pages:
                pages[page_num] = []
            pages[page_num].append(element)
        
        # Process each page
        for page_num in sorted(pages.keys()):
            page_elements = pages[page_num]
            page_order = self._analyze_page_reading_order(page_elements)
            reading_order.extend(page_order)
        
        return reading_order
    
    def _analyze_page_reading_order(self, elements: List[Dict]) -> List[str]:
        """Analyze reading order for a single page"""
        order = []
        
        # Sort elements by position (top to bottom, then right to left for Hebrew)
        positioned_elements = [e for e in elements if e.get('bbox')]
        
        if self.hebrew_aware:
            # Hebrew reading order: top to bottom, right to left
            positioned_elements.sort(key=lambda x: (x['bbox'][1], -x['bbox'][0]))
        else:
            # Standard reading order: top to bottom, left to right
            positioned_elements.sort(key=lambda x: (x['bbox'][1], x['bbox'][0]))
        
        for element in positioned_elements:
            element_id = element.get('chunk_id', f"element_{id(element)}")
            order.append(element_id)
        
        return order
    
    def identify_layout_structure(self, elements: List[Dict]) -> Dict[str, Any]:
        """Identify layout structure (columns, headers, etc.)"""
        structure = {
            "columns": [],
            "headers": [],
            "sections": []
        }
        
        # Simple column detection based on x-coordinates
        positioned_elements = [e for e in elements if e.get('bbox')]
        
        if positioned_elements:
            # Detect columns by clustering x-coordinates
            x_coords = [e['bbox'][0] for e in positioned_elements]
            columns = self._detect_columns(x_coords)
            structure["columns"] = columns
            
            # Detect headers (elements with larger font or special positioning)
            headers = self._detect_headers(positioned_elements)
            structure["headers"] = headers
        
        return structure
    
    def _detect_columns(self, x_coords: List[float]) -> List[Dict]:
        """Detect column structure from x-coordinates"""
        if not x_coords:
            return []
        
        # Simple clustering based on x-coordinate gaps
        sorted_x = sorted(set(x_coords))
        columns = []
        current_col_start = sorted_x[0]
        
        for i in range(1, len(sorted_x)):
            gap = sorted_x[i] - sorted_x[i-1]
            if gap > 50:  # Significant gap indicates new column
                columns.append({
                    "start_x": current_col_start,
                    "end_x": sorted_x[i-1],
                    "column_index": len(columns)
                })
                current_col_start = sorted_x[i]
        
        # Add last column
        columns.append({
            "start_x": current_col_start,
            "end_x": sorted_x[-1],
            "column_index": len(columns)
        })
        
        return columns
    
    def _detect_headers(self, elements: List[Dict]) -> List[str]:
        """Detect header elements"""
        headers = []
        
        for element in elements:
            content = element.get('content', '')
            bbox = element.get('bbox')
            
            # Simple heuristics for header detection
            if (len(content) < 100 and  # Short text
                bbox and bbox[1] < 100):  # Near top of page
                element_id = element.get('chunk_id', f"element_{id(element)}")
                headers.append(element_id)
        
        return headers
    
    def find_element_relationships(self, elements: List[Dict]) -> Dict[str, List[str]]:
        """Find spatial relationships between elements"""
        relationships = {}
        
        for i, element in enumerate(elements):
            element_id = element.get('chunk_id', f"element_{i}")
            relationships[element_id] = {
                "above": [],
                "below": [],
                "left": [],
                "right": [],
                "nearby": []
            }
            
            element_bbox = element.get('bbox')
            if not element_bbox:
                continue
            
            # Compare with other elements
            for j, other_element in enumerate(elements):
                if i == j:
                    continue
                
                other_bbox = other_element.get('bbox')
                if not other_bbox:
                    continue
                
                other_id = other_element.get('chunk_id', f"element_{j}")
                relationship = self._determine_spatial_relationship(element_bbox, other_bbox)
                
                if relationship:
                    relationships[element_id][relationship].append(other_id)
        
        return relationships
    
    def _determine_spatial_relationship(self, bbox1: tuple, bbox2: tuple) -> Optional[str]:
        """Determine spatial relationship between two bounding boxes"""
        # Calculate centers
        center1 = ((bbox1[0] + bbox1[2]) / 2, (bbox1[1] + bbox1[3]) / 2)
        center2 = ((bbox2[0] + bbox2[2]) / 2, (bbox2[1] + bbox2[3]) / 2)
        
        # Thresholds for determining relationships
        horizontal_threshold = 20
        vertical_threshold = 20
        
        # Vertical relationships
        if center2[1] < center1[1] - vertical_threshold:
            return "above"
        elif center2[1] > center1[1] + vertical_threshold:
            return "below"
        
        # Horizontal relationships (for elements on same level)
        elif abs(center2[1] - center1[1]) <= vertical_threshold:
            if center2[0] < center1[0] - horizontal_threshold:
                return "left"
            elif center2[0] > center1[0] + horizontal_threshold:
                return "right"
        
        # Check if nearby
        distance = ((center1[0] - center2[0]) ** 2 + (center1[1] - center2[1]) ** 2) ** 0.5
        if distance < 100:  # Within 100 pixels
            return "nearby"
        
        return None